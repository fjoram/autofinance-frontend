from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 0:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
    return p

def add_meta(label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label} ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_value = p.add_run(value)
    run_value.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table_from_rows(rows):
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text.strip('* ')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else p.add_run(cell.text)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                set_cell_bg(cell, '1A3A6B')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:
                set_cell_bg(cell, 'EEF2FF')
    doc.add_paragraph()

with open('BUSINESS_PROPOSAL.md', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # Title (# H1)
    if line.startswith('# ') and not line.startswith('## '):
        add_heading(line[2:], 0)
        i += 1
        continue

    # H2
    if line.startswith('## '):
        doc.add_paragraph()
        add_heading(line[3:], 1)
        i += 1
        continue

    # H3
    if line.startswith('### '):
        add_heading(line[4:], 2)
        i += 1
        continue

    # Horizontal rule
    if line.startswith('---'):
        i += 1
        continue

    # Meta lines (bold key: value)
    if line.startswith('**') and ':**' in line:
        parts = line.split(':**', 1)
        label = parts[0].replace('**', '').strip() + ':'
        value = parts[1].replace('**', '').strip()
        add_meta(label, value)
        i += 1
        continue

    # Table
    if line.startswith('|'):
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row_line = lines[i].strip()
            if re.match(r'^\|[-| :]+\|$', row_line):
                i += 1
                continue
            cells = [c.strip() for c in row_line.strip('|').split('|')]
            table_rows.append(cells)
            i += 1
        if table_rows:
            add_table_from_rows(table_rows)
        continue

    # Bullet points
    if line.startswith('- '):
        text = line[2:]
        p = doc.add_paragraph(style='List Bullet')
        # Handle inline bold
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(11)
            if idx % 2 == 1:
                run.bold = True
        i += 1
        continue

    # Italic/bold result line
    if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line.strip('*'))
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
        i += 1
        continue

    # Bold result line **text**
    if line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line.strip('*'))
        run.bold = True
        run.font.size = Pt(11)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        p.add_run(m.group(2)).font.size = Pt(11)
        i += 1
        continue

    # Empty line
    if line == '':
        i += 1
        continue

    # Normal paragraph with possible inline bold
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.+?)\*\*', line)
    for idx, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(11)
        if idx % 2 == 1:
            run.bold = True
    i += 1

doc.save('AutoFinance_Hub_Business_Proposal_IM_Bank.docx')
print("Saved: AutoFinance_Hub_Business_Proposal_IM_Bank.docx")
